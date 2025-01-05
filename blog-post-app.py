import streamlit as st
import openai
from docx import Document  # For Word file generation
from fpdf import FPDF  # For PDF file generation
import os
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Use environment variable for API key
openai.api_key = os.getenv("OPENAI_API_KEY")



# Helper function: Convert Word document to bytes
def doc_to_bytes(doc):
    """Convert a Word document to bytes."""
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# App Title
st.title("Pavelist Blog Generator")

# User Inputs
st.subheader("Let’s Create a Unique Blog Post Tailored to Your Needs")
blog_subject = st.text_input("Enter the Blog Topic:", "")
keywords = st.text_input("Enter Keywords (comma-separated):", "")
language_level = st.selectbox("Select English Level:", ["Basic", "Intermediate", "Advanced"])
word_count = st.number_input("Enter Number of Words:", min_value=50, max_value=2000, step=50, value=500)
additional_details = st.text_area("Provide Additional Instructions (Optional):", "")

# Checkbox for Human-Like Writing
human_like = st.checkbox("Make it Sound Like Human Writing")

# Generate Button
if st.button("Generate "):
    if blog_subject and word_count > 0:
        # Format the prompt
        prompt = f"""
        Write a blog post about: {blog_subject}.
        Use the following keywords: {keywords}.
        The blog post should be approximately {word_count} words long.
        {additional_details}
        {"Make it feel like a human wrote it, with natural language, conversational tone, and subtle imperfections." if human_like else ""}
        """
        try:
            # Call OpenAI API
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=min(4000, word_count * 4),  # Adjust max tokens
                temperature=0.85 if human_like else 0.7
            )
            blog_post = response.choices[0].message["content"]
            
            # Display the blog post
            st.subheader("Generated Blog Post:")
            st.write(blog_post)
            
            # Sanitize the blog title for file naming
            sanitized_title = "".join(
                char if char.isalnum() or char in (" ", "-", "_") else "_" for char in blog_subject.strip()
            ).replace(" ", "_")
            
            # Word File Download
            doc = Document()
            doc.add_heading(blog_subject, level=1)
            doc.add_paragraph(blog_post)
            word_file_name = f"{sanitized_title}.docx"
            st.download_button(
                label="Download as Word Document",
                data=doc_to_bytes(doc),
                file_name=word_file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # PDF File Download
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            # Handle encoding issues by replacing unsupported characters
            clean_text = blog_post.encode("ascii", "ignore").decode("ascii")
            pdf.multi_cell(0, 10, clean_text)
            pdf_file_name = f"{sanitized_title}.pdf"
            pdf_data = pdf.output(dest='S').encode('latin1')
            st.download_button(
                label="Download as PDF",
                data=pdf_data,
                file_name=pdf_file_name,
                mime="application/pdf"
            )
        except openai.error.OpenAIError as e:
            st.error(f"OpenAI API error: {e}")
    else:
        st.warning("Please provide a blog topic and valid word count to generate content.")
